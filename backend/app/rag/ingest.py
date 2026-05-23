import os
from pathlib import Path
from typing import List

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document

from .embeddings import get_embedding_function


CHROMA_PERSIST_DIR = os.getenv("CHROMA_PERSIST_DIR", "./data/chroma")
COLLECTION_NAME = "research_docs"


def load_pdf(file_path: str) -> List[Document]:
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    documents = []
    for page_num, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            documents.append(
                Document(
                    page_content=text,
                    metadata={"source": file_path, "page": page_num + 1},
                )
            )
    doc.close()
    return documents


def load_docx(file_path: str) -> List[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [Document(page_content=full_text, metadata={"source": file_path})]


def load_document(file_path: str) -> List[Document]:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return load_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def ingest_document(file_path: str) -> int:
    """
    Ingest a document into ChromaDB.
    Returns the number of chunks stored.
    """
    raw_docs = load_document(file_path)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", ".", " ", ""],
    )
    chunks = splitter.split_documents(raw_docs)

    if not chunks:
        return 0

    embedding_fn = get_embedding_function()

    vectorstore = Chroma(
        collection_name=COLLECTION_NAME,
        embedding_function=embedding_fn,
        persist_directory=CHROMA_PERSIST_DIR,
    )
    vectorstore.add_documents(chunks)

    return len(chunks)